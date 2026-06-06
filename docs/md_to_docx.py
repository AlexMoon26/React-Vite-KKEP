from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DEFAULT_INPUT = ROOT / "01-react-vite-architecture.md"


def add_inline_code(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size in [("Heading 1", 20), ("Heading 2", 15), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(32, 32, 32)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_code_block(doc, lines, language):
    if language:
        label = doc.add_paragraph()
        label.paragraph_format.space_before = Pt(3)
        label.paragraph_format.space_after = Pt(1)
        run = label.add_run(language)
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)

    text = "\n".join(lines) if lines else ""
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(35, 35, 35)


def parse_table(lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for cell_index, value in enumerate(row):
            p = cells[cell_index].paragraphs[0]
            add_inline_code(p, value)
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def add_paragraph_with_inline(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    add_inline_code(paragraph, text)


def convert():
    input_path = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else DEFAULT_INPUT
    output_path = input_path.with_suffix(".docx")

    doc = Document()
    style_document(doc)

    lines = input_path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_language = ""
    code_lines = []
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_table(doc, parse_table(table_lines))
            table_lines = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines, code_language)
                in_code = False
                code_language = ""
                code_lines = []
            else:
                flush_table()
                in_code = True
                code_language = line[3:].strip()
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            continue
        flush_table()

        if not line.strip():
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            add_paragraph_with_inline(doc, line[2:].strip(), style="List Bullet")
            continue
        if re.match(r"^\d+\. ", line):
            add_paragraph_with_inline(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
            continue

        add_paragraph_with_inline(doc, line)

    flush_table()

    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    try:
        convert()
    except Exception as exc:
        print(f"Failed to convert Markdown to DOCX: {exc}", file=sys.stderr)
        raise
